#!/usr/bin/env python3
"""fill_forms.py — r08youshiki1_5.docx テーブルフォーム記入スクリプト

data/source/r08youshiki1_5.docx をコピーし、YAML データに基づいて
各テーブル様式にデータを書き込む。不要な様式は削除する。

Usage:
    python fill_forms.py \\
        --config main/00_setup/config.yaml \\
        --researchers main/00_setup/researchers.yaml \\
        --other-funding main/00_setup/other_funding.yaml \\
        --source data/source/r08youshiki1_5.docx \\
        --output main/step02_docx/output/
"""

import argparse
import copy
import math
import re
import shutil
import sys
import warnings
from datetime import date as _date
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table as DocxTable


# ============================================================================
# YAML helpers
# ============================================================================

def load_yaml(path):
    """Load YAML file."""
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


# ============================================================================
# Cell-level helpers
# ============================================================================

def _first_rpr(cell):
    """Return deepcopy of first run's rPr, or None."""
    for p in cell.paragraphs:
        for r in p.runs:
            rpr = r._element.find(qn("w:rPr"))
            if rpr is not None:
                return copy.deepcopy(rpr)
    return None


def _first_ppr(cell):
    """Return deepcopy of first paragraph's pPr, or None."""
    if cell.paragraphs:
        ppr = cell.paragraphs[0]._element.find(qn("w:pPr"))
        if ppr is not None:
            return copy.deepcopy(ppr)
    return None


def set_cell(cell, text):
    r"""Replace all cell content.  ``\n`` → new paragraph.  Preserves first-run format."""
    text = str(text) if text is not None else ""
    rpr = _first_rpr(cell)
    ppr = _first_ppr(cell)
    tc = cell._element
    for old_p in list(tc.findall(qn("w:p"))):
        tc.remove(old_p)
    for line in (text.split("\n") if text else [""]):
        p = OxmlElement("w:p")
        if ppr is not None:
            p.append(copy.deepcopy(ppr))
        r = OxmlElement("w:r")
        if rpr is not None:
            r.append(copy.deepcopy(rpr))
        t = OxmlElement("w:t")
        t.text = line
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)


def circle_choice(cell, target):
    """Mark *target* with ○ in cell to indicate selection."""
    for p in cell.paragraphs:
        for run in p.runs:
            if target in run.text:
                run.text = run.text.replace(target, f"○{target}")
                return True
    # fallback: work with assembled text
    full = cell.text
    if target in full:
        set_cell(cell, full.replace(target, f"○{target}"))
        return True
    warnings.warn(f"circle_choice: '{target}' not found in '{full[:60]}'")
    return False


def _ensure_table_rows(tbl, target_count, template_row_idx=None):
    """テーブルに target_count 行以上が存在することを保証する。

    不足分は template_row_idx を deep copy して追加する。set_cell は段落を
    完全置換するため、複製先の文字列が残ることはない。罫線・セル幅・段落
    プロパティが保持されることで、Word 上の見た目が崩れない。

    Args:
        tbl: docx.table.Table
        target_count: 必要な行数
        template_row_idx: コピー元の行インデックス。None なら最終行を使用
    """
    while len(tbl.rows) < target_count:
        src_idx = template_row_idx if template_row_idx is not None else len(tbl.rows) - 1
        src_tr = tbl.rows[src_idx]._tr
        new_tr = copy.deepcopy(src_tr)
        tbl._tbl.append(new_tr)


# ============================================================================
# Format helpers
# ============================================================================

def _fw(s):
    """ASCII → full-width."""
    return "".join(
        chr(ord(c) + 0xFEE0) if 0x21 <= ord(c) <= 0x7E
        else ("\u3000" if c == " " else c)
        for c in str(s)
    )


def _amt(n):
    """Comma-formatted integer string."""
    return f"{int(n):,}" if isinstance(n, (int, float)) else str(n)


def _budget_year(yr, rate):
    """Calculate derived budget fields for one fiscal year."""
    e = yr.get("equipment", 0)
    c = yr.get("consumables", 0)
    t = yr.get("travel", 0)
    p = yr.get("personnel", 0)
    o = yr.get("other", 0)
    d = e + c + t + p + o
    i = math.ceil(d * rate)
    return dict(goods=e + c, travel=t, personnel=p, other=o,
                direct=d, indirect=i, total=d + i)


# ============================================================================
# Table identification
# ============================================================================

def identify_tables(doc):
    """Pattern-match each form table.  Returns ``{key: Table}``."""
    found = {}
    detail = []   # 様式2-2
    cv = []        # 様式4
    form3 = []     # 様式3

    for table in doc.tables:
        nr, nc = len(table.rows), len(table.columns)
        c00 = table.cell(0, 0).text

        # 様式1-1: 20r×8c  ①研究テーマ in top-left
        if "①研究テーマ" in c00 and nr >= 20:
            found["1-1"] = table
            continue

        # 様式2-1 費目内訳: 9r×7c  contains ア．物品費
        if nr == 9 and nc == 7:
            if any("ア．物品費" in table.cell(r, 0).text for r in range(nr)):
                found["2-1"] = table
                continue

        # 様式2-1 機関別: 6r×7c  header matches + appears after 2-1
        if nr == 6 and nc == 7 and "研究費の内訳" in c00 and "2-1" in found:
            found["2-1_inst"] = table
            continue

        # 様式2-2: 12r×5c  Ⅰ in row-1 col-0
        if nr == 12 and nc == 5:
            r1 = table.cell(1, 0).text
            if "Ⅰ" in r1 or "物品費" in r1:
                detail.append(table)
                continue

        # 様式3: 6r×8c  【本研究課題】 in row-1 col-2
        if nr >= 5 and nc == 8:
            try:
                if any("【本研究課題】" in table.cell(r, 2).text
                       for r in range(min(nr, 3))):
                    form3.append(table)
                    continue
            except Exception:
                pass

        # 様式4: 10r×5c  研究課題名 in top-left
        if nr == 10 and nc == 5 and "研究課題名" in c00:
            cv.append(table)
            continue

        # 様式5: 26r×7c  企業名
        if nr >= 25 and nc == 7 and "企" in c00:
            found["5"] = table
            continue

        # チェックリスト: Nx2  チェック in header
        if nc == 2 and any("チェック" in c.text for c in table.rows[0].cells):
            found["checklist"] = table
            continue

    for i, t in enumerate(detail):
        found[f"2-2_y{i + 1}"] = t
    if form3:
        found["3-1"] = form3[0]
    if len(form3) > 1:
        found["3-2"] = form3[1]
    for i, t in enumerate(cv):
        found[f"4-{i + 1}"] = t
    return found


# ============================================================================
# Fill: 様式1-1  (Table 0: 20r×8c)
# ============================================================================

def fill_1_1(tbl, cfg, res, ofund):
    """Fill 様式1-1 申請書概要."""
    proj = cfg["project"]
    pi = res["pi"]
    admin = res["admin_contact"]
    inst = cfg["lead_institution"]
    budget = cfg["budget"]
    rate = budget.get("indirect_rate", 0.3)

    # ① 研究テーマ
    set_cell(tbl.cell(0, 3), str(proj["theme_number"]))

    # ② 研究課題名
    set_cell(tbl.cell(1, 3), proj["title_ja"])
    set_cell(tbl.cell(2, 3), proj["title_en"])

    # ③ 研究分野 — ○ で選択
    circle_choice(tbl.cell(3, 3), proj["field"])

    # ④ キーワード
    kw = proj.get("keywords", [])
    set_cell(tbl.cell(4, 3), "、".join(str(k) for k in kw) if kw else "")

    # ⑤ 研究の概要
    set_cell(tbl.cell(5, 3), proj.get("summary", ""))

    # ⑥ 研究期間
    s = str(proj.get("period_start", "R8")).replace("R", "")
    e = str(proj.get("period_end", "R10")).replace("R", "")
    y = str(proj.get("period_years", 3))
    set_cell(tbl.cell(6, 3),
             f"令和{_fw(s)}年度 ～ 令和{_fw(e)}年度（{_fw(y)}か年度）")

    # ⑦ 申請額
    total = sum(_budget_year(yr, rate)["total"] for yr in budget["yearly"])
    set_cell(tbl.cell(7, 3), f"{_amt(total)}千円")

    # ⑧ タイプ — ○ で選択
    tp = str(proj.get("type", "A"))
    circle_choice(tbl.cell(8, 3), f"タイプ{_fw(tp)}")

    # ⑧ 重複応募
    if proj.get("duplicate_application"):
        circle_choice(tbl.cell(8, 7), "有")
    else:
        circle_choice(tbl.cell(8, 7), "無")

    # ⑨ 研究代表者 (rows 9–11)
    set_cell(tbl.cell(9, 3), pi["name_ja"])
    set_cell(tbl.cell(9, 7), pi.get("nationality", ""))

    set_cell(tbl.cell(10, 3),
             f"{pi.get('affiliation', '')}・"
             f"{pi.get('department', '')}・"
             f"{pi.get('position', '')}")

    ct = pi.get("contact", {})
    postal = ct.get("postal", "").replace("〒", "").strip()
    set_cell(tbl.cell(11, 3),
             f"〒{postal}\nTEL: {ct.get('tel', '')}\nE-mail: {ct.get('email', '')}")

    # ⑩ 経理事務担当者 (rows 12–14)
    set_cell(tbl.cell(12, 3), admin.get("name", ""))

    set_cell(tbl.cell(13, 3),
             f"{admin.get('affiliation', '')}・"
             f"{admin.get('department', '')}・"
             f"{admin.get('position', '')}")

    set_cell(tbl.cell(14, 3),
             f"〒\nTEL: {admin.get('tel', '')}\nE-mail: {admin.get('email', '')}")

    # ⑪ 研究者リスト (rows 17–)
    # Dynamically add rows if template doesn't have enough for all co-investigators
    co_list = res.get("co_investigators", [])
    needed_rows = 18 + len(co_list)
    while len(tbl.rows) < needed_rows:
        src_tr = tbl.rows[-1]._element
        new_tr = copy.deepcopy(src_tr)
        tbl._element.append(new_tr)

    set_cell(tbl.cell(17, 0), inst.get("name", ""))
    set_cell(tbl.cell(17, 3), f"研究代表者\n{pi['name_ja']}")
    set_cell(tbl.cell(17, 4),
             f"{pi.get('department', '')}・{pi.get('position', '')}\n"
             f"TEL: {ct.get('tel', '')}\nE-mail: {ct.get('email', '')}")

    for idx, co in enumerate(co_list):
        row = 18 + idx
        co_ct = co.get("contact", {})
        set_cell(tbl.cell(row, 0),
                 co.get("institution", co.get("affiliation", "")))
        set_cell(tbl.cell(row, 3), f"研究分担者\n{co['name_ja']}")
        set_cell(tbl.cell(row, 4),
                 f"{co.get('department', '')}・{co.get('position', '')}\n"
                 f"TEL: {co_ct.get('tel', '')}\n"
                 f"E-mail: {co_ct.get('email', '')}")

    print("  ✓ 様式1-1")


# ============================================================================
# Fill: 様式2-1 費目内訳  (Table 3: 9r×7c)
# ============================================================================

def fill_2_1(tbl, cfg):
    """Fill 様式2-1 年度別費目内訳."""
    budget = cfg["budget"]
    rate = budget.get("indirect_rate", 0.3)
    for yr in budget["yearly"]:
        col = yr["year"]
        if col > 5:
            continue
        b = _budget_year(yr, rate)
        set_cell(tbl.cell(2, col), _amt(b["direct"]))     # 直接経費
        set_cell(tbl.cell(3, col), _amt(b["goods"]))       # ア．物品費
        set_cell(tbl.cell(4, col), _amt(b["personnel"]))   # イ．人件費・謝金
        set_cell(tbl.cell(5, col), _amt(b["travel"]))      # ウ．旅費
        set_cell(tbl.cell(6, col), _amt(b["other"]))       # エ．その他
        set_cell(tbl.cell(7, col), _amt(b["indirect"]))    # 間接経費
        set_cell(tbl.cell(8, col), _amt(b["total"]))       # 合計
    print("  ✓ 様式2-1 費目内訳")


# ============================================================================
# Fill: 様式2-1 機関別  (Table 4: 6r×7c)
# ============================================================================

def fill_2_1_inst(tbl, cfg):
    """Fill 様式2-1 機関別研究費."""
    budget = cfg["budget"]
    rate = budget.get("indirect_rate", 0.3)
    by_inst = budget.get("by_institution", [])

    for ii, inst_data in enumerate(by_inst):
        row = 2 + ii
        if row >= len(tbl.rows) - 1:
            warnings.warn("様式2-1機関別: not enough rows")
            break
        set_cell(tbl.cell(row, 0), inst_data["institution"])
        for yr in inst_data.get("yearly", []):
            col = yr["year"]
            if col > 5:
                continue
            amt = yr.get("amount", 0)
            ind = math.ceil(amt * rate)
            tot = amt + ind
            set_cell(tbl.cell(row, col), f"{_amt(tot)}\n({_amt(ind)})")

    # totals row (last row)
    last = len(tbl.rows) - 1
    for yr in budget["yearly"]:
        col = yr["year"]
        if col > 5:
            continue
        b = _budget_year(yr, rate)
        set_cell(tbl.cell(last, col), f"{_amt(b['total'])}\n({_amt(b['indirect'])})")
    print("  ✓ 様式2-1 機関別")


# ============================================================================
# Fill: 様式2-2  (Tables 5–9: 12r×5c each)
# ============================================================================

# N16-02: 様式2-2 のセル内テキストはデフォルトの明朝＋10pt だと 5 列構成
# （内訳/数量/金額/機関/根拠）の各セルが狭く、長い品目名や根拠説明が極端に
# 折り返される。共同研究者指摘 #3 に従い、プロポーショナルゴシック・9pt に
# 揃えることで可読性を確保する。reference.docx ではなく run 単位で rPr を
# 上書きするため、明朝指定の Normal スタイルから影響を受けない。
#
# N16-03: ＭＳ ゴシック（等幅）だと半角数字も全幅扱いとなり、金額列の
# 「15,000」(6 文字) が約 1080twips に膨らんで列幅 1078twips を超え、
# Word が「15,0/00」のようにセル内で折り返す。プロポーショナル変種
# ＭＳ Ｐゴシックは半角数字を約 0.5em 幅で描画するため、「15,000」も
# 約 540twips に収まり折り返さない。
_GOTHIC_FONT = "ＭＳ Ｐゴシック"


def _apply_gothic_to_cell(cell, sz_half_pt=18):
    """Force all runs in a cell to MS Gothic (sz_half_pt half-points = 9pt).

    N16-06: テンプレート（r08youshiki1_5.docx）の様式2-2 金額セルの段落には
    `<w:ind w:right="732"/>` が設定されており、列幅 1276twips のうち 732twips
    が右インデントで予約されてしまい、有効幅 346twips しか残らない。これでは
    「1,000」「2,000」が「1,00 / 0」のように途中改行される。フォント設定と
    一緒に右インデントもリセットすることで本来の列幅を活用させる。
    """
    for p in cell.paragraphs:
        # Remove right indent on paragraph (preserve other ind attrs)
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            ind = pPr.find(qn("w:ind"))
            if ind is not None and ind.get(qn("w:right")) is not None:
                ind.set(qn("w:right"), "0")
        for r in p.runs:
            rPr = r._element.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r._element.insert(0, rPr)
            for old in rPr.findall(qn("w:rFonts")):
                rPr.remove(old)
            for old in rPr.findall(qn("w:sz")):
                rPr.remove(old)
            for old in rPr.findall(qn("w:szCs")):
                rPr.remove(old)
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), _GOTHIC_FONT)
            rFonts.set(qn("w:eastAsia"), _GOTHIC_FONT)
            rFonts.set(qn("w:hAnsi"), _GOTHIC_FONT)
            rFonts.set(qn("w:cs"), _GOTHIC_FONT)
            rPr.insert(0, rFonts)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(sz_half_pt))
            rPr.append(sz)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), str(sz_half_pt))
            rPr.append(szCs)


def _fill_budget_block(tbl, row, sections):
    """Write aligned multi-line content into a budget-detail row.

    *sections*: list of ``(header_str | None, items_list | None)``
    """
    cols = [[] for _ in range(5)]
    for hdr, items in sections:
        if hdr:
            cols[0].append(hdr)
            for c in range(1, 5):
                cols[c].append("")
        if items:
            for it in items:
                cols[0].append(it.get("name", ""))
                cols[1].append(it.get("quantity", ""))
                cols[2].append(_amt(it.get("amount", 0)))
                cols[3].append(it.get("institution", ""))
                cols[4].append(it.get("justification", ""))
    for c in range(5):
        set_cell(tbl.cell(row, c), "\n".join(cols[c]))


def fill_2_2(tables, cfg):
    """Fill 様式2-2 年度別研究費計画書."""
    budget = cfg["budget"]
    rate = budget.get("indirect_rate", 0.3)

    for yd in budget.get("details", []):
        yr = yd["year"]
        key = f"2-2_y{yr}"
        if key not in tables:
            warnings.warn(f"様式2-2 year {yr} table not found")
            continue
        tbl = tables[key]
        items = yd.get("line_items", [])
        equip = [i for i in items if i["category"] == "equipment"]
        consu = [i for i in items if i["category"] == "consumables"]
        pers  = [i for i in items if i["category"] == "personnel"]
        trav  = [i for i in items if i["category"] == "travel"]
        othr  = [i for i in items if i["category"] == "other"]

        # Row 1: Ⅰ．物品費  (equipment + consumables)
        _fill_budget_block(tbl, 1, [
            ("Ⅰ．物品費", None),
            ("1.設備備品費", equip),
            ("2.消耗品費", consu),
        ])
        g_tot = sum(i.get("amount", 0) for i in equip + consu)
        set_cell(tbl.cell(2, 2), _amt(g_tot))   # 小計

        # Row 3: Ⅱ．人件費・謝金
        _fill_budget_block(tbl, 3, [
            ("Ⅱ．人件費・謝金", None),
            ("1.人件費", pers),
        ])
        p_tot = sum(i.get("amount", 0) for i in pers)
        set_cell(tbl.cell(4, 2), _amt(p_tot))

        # Row 5: Ⅲ．旅費
        _fill_budget_block(tbl, 5, [
            ("Ⅲ．旅費", None),
            ("1.旅費", trav),
        ])
        t_tot = sum(i.get("amount", 0) for i in trav)
        set_cell(tbl.cell(6, 2), _amt(t_tot))

        # Row 7: Ⅳ．その他
        _fill_budget_block(tbl, 7, [
            ("Ⅳ．その他", None),
            (None, othr),
        ])
        o_tot = sum(i.get("amount", 0) for i in othr)
        set_cell(tbl.cell(8, 2), _amt(o_tot))

        # Totals
        direct = g_tot + p_tot + t_tot + o_tot
        indirect = math.ceil(direct * rate)
        set_cell(tbl.cell(9, 2), _amt(direct))
        set_cell(tbl.cell(10, 2), _amt(indirect))
        set_cell(tbl.cell(11, 2), _amt(direct + indirect))

        # N16-02: ゴシック化＋9pt で可読性を確保（共同研究者指摘 #3）
        for row in tbl.rows:
            for cell in row.cells:
                _apply_gothic_to_cell(cell)

    print("  ✓ 様式2-2")


# ============================================================================
# Fill: 様式3-1  (Table 10: 6r×8c)
# ============================================================================

def fill_3_1(tbl, cfg, res, ofund):
    """Fill 様式3-1 他制度応募状況（研究代表者）."""
    proj = cfg["project"]
    pi = res["pi"]
    budget = cfg["budget"]
    rate = budget.get("indirect_rate", 0.3)

    start = str(proj.get("period_start", "R8"))
    end = str(proj.get("period_end", "R10"))

    # Row 1: 本研究課題
    set_cell(tbl.cell(1, 2), f"【本研究課題】\n{start}～{end}\n防衛装備庁")
    set_cell(tbl.cell(1, 3), proj["title_ja"])
    set_cell(tbl.cell(1, 4), "代表")

    totals = [_budget_year(yr, rate)["total"] for yr in budget["yearly"]]
    yr1 = totals[0] if totals else 0
    grand = sum(totals)
    set_cell(tbl.cell(1, 5), f"{_amt(yr1)}\n({_amt(grand)})")
    set_cell(tbl.cell(1, 6), str(pi.get("effort_percent", "")))

    # Other funding entries
    entries = ofund.get("pi_funding", {}).get("entries", [])
    if not entries:
        set_cell(tbl.cell(2, 2), "無し")
    else:
        # entries 数に応じてテーブル行を動的拡張（最初のデータ行 2 をテンプレートに）
        _ensure_table_rows(tbl, target_count=2 + len(entries), template_row_idx=2)
        for idx, ent in enumerate(entries):
            row = 2 + idx
            set_cell(tbl.cell(row, 0), str(idx + 2))
            set_cell(tbl.cell(row, 1), ent.get("status", ""))
            agency = ent.get("agency", "") if not ent.get("confidential") else ""
            set_cell(tbl.cell(row, 2),
                     f"{ent.get('program_name', '')}\n"
                     f"{ent.get('period', '')}\n{agency}")
            set_cell(tbl.cell(row, 3), ent.get("project_title", ""))
            set_cell(tbl.cell(row, 4), ent.get("role", ""))
            if ent.get("confidential"):
                set_cell(tbl.cell(row, 5), "")
            else:
                set_cell(tbl.cell(row, 5),
                         f"{_amt(ent.get('budget', 0))}\n"
                         f"({_amt(ent.get('total_budget', 0))})")
            set_cell(tbl.cell(row, 6), str(ent.get("effort_percent", "")))
            set_cell(tbl.cell(row, 7), ent.get("difference", ""))

    print("  ✓ 様式3-1")


# ============================================================================
# Fill: 様式3-2  (Table 11: 6r×8c)
# ============================================================================

def fill_3_2(tbl, cfg, res, ofund):
    """Fill 様式3-2 他制度応募状況（研究分担者）."""
    proj = cfg["project"]
    budget = cfg["budget"]
    rate = budget.get("indirect_rate", 0.3)

    start = str(proj.get("period_start", "R8"))
    end = str(proj.get("period_end", "R10"))

    # Row 1: 本研究課題
    set_cell(tbl.cell(1, 2), f"【本研究課題】\n{start}～{end}\n防衛装備庁")
    set_cell(tbl.cell(1, 3), proj["title_ja"])
    set_cell(tbl.cell(1, 4), "分担")

    totals = [_budget_year(yr, rate)["total"] for yr in budget["yearly"]]
    yr1 = totals[0] if totals else 0
    grand = sum(totals)
    set_cell(tbl.cell(1, 5), f"{_amt(yr1)}\n({_amt(grand)})")

    co_invs = res.get("co_investigators", [])
    efforts = [str(co.get("effort_percent", "")) for co in co_invs]
    set_cell(tbl.cell(1, 6), ", ".join(efforts) if efforts else "")

    # Collect all entries across all co-investigators
    all_entries = []
    for co_fund in ofund.get("co_investigator_funding", []):
        rname = co_fund.get("researcher_name", "")
        for ent in co_fund.get("entries", []):
            all_entries.append((rname, ent))

    if not all_entries:
        set_cell(tbl.cell(2, 2), "無し")
    else:
        # entries 数に応じてテーブル行を動的拡張（最初のデータ行 2 をテンプレートに）
        _ensure_table_rows(tbl, target_count=2 + len(all_entries), template_row_idx=2)
        for idx, (rname, ent) in enumerate(all_entries):
            row = 2 + idx
            set_cell(tbl.cell(row, 0), str(idx + 2))
            set_cell(tbl.cell(row, 1), ent.get("status", ""))
            agency = ent.get("agency", "") if not ent.get("confidential") else ""
            set_cell(tbl.cell(row, 2),
                     f"{ent.get('program_name', '')}\n"
                     f"{ent.get('period', '')}\n{agency}")
            set_cell(tbl.cell(row, 3),
                     f"{ent.get('project_title', '')}\n（{rname}）")
            set_cell(tbl.cell(row, 4), ent.get("role", ""))
            if ent.get("confidential"):
                set_cell(tbl.cell(row, 5), "")
            else:
                set_cell(tbl.cell(row, 5),
                         f"{_amt(ent.get('budget', 0))}\n"
                         f"({_amt(ent.get('total_budget', 0))})")
            set_cell(tbl.cell(row, 6), str(ent.get("effort_percent", "")))
            set_cell(tbl.cell(row, 7), ent.get("difference", ""))

    print("  ✓ 様式3-2")


# ============================================================================
# Fill: 様式4  (Tables 12–13: 10r×5c)
# ============================================================================

def fill_4(tbl, cfg, person, label="様式4-1"):
    """Fill 様式4-1 or 4-2 研究者調書."""
    proj = cfg["project"]

    # Row 0: 研究課題名
    set_cell(tbl.cell(0, 2), proj["title_ja"])

    # Row 1: ふりがな/氏名, 生年月日/年齢
    set_cell(tbl.cell(1, 2),
             f"{person.get('furigana', '')}\n{person['name_ja']}")
    age = person.get("age", "")
    set_cell(tbl.cell(1, 4),
             f"{person.get('birth_date', '')}\n（{age}歳）")

    # Row 2: 研究者番号
    set_cell(tbl.cell(2, 2), person.get("researcher_id", ""))

    # Row 2–3, col 4: 最終学歴・学位 (vertically merged)
    edu = person.get("education", {})
    set_cell(tbl.cell(2, 4),
             f"{edu.get('school', '')}\n"
             f"{edu.get('graduation_year', '')}\n"
             f"{edu.get('degree', '')}")

    # Row 3: 所属機関・部局・職/職階
    set_cell(tbl.cell(3, 2),
             f"{person.get('affiliation', '')}・"
             f"{person.get('department', '')}・"
             f"{person.get('position', '')}")

    # Row 4: 専門分野
    set_cell(tbl.cell(4, 2), person.get("specialty", ""))

    # Row 5: 主な経歴
    career = person.get("career_history", [])
    set_cell(tbl.cell(5, 2),
             "\n".join(f"{c.get('year_month', '')}\t{c.get('description', '')}"
                       for c in career) if career else "")

    # Row 6: 競争的研究資金獲得実績
    funding = person.get("funding_history", [])
    lines = []
    for f in funding:
        lines.append(
            f"{f.get('program', '')} 「{f.get('title', '')}」 "
            f"{f.get('role', '')} {f.get('period', '')} {f.get('amount', '')}")
    set_cell(tbl.cell(6, 2), "\n".join(lines) if lines else "")

    # Row 7: 受賞歴・表彰歴
    awards = person.get("awards", [])
    if awards:
        set_cell(tbl.cell(7, 2),
                 "\n".join(f"{a.get('year_month', '')}\t{a.get('name', '')}"
                           for a in awards))
    else:
        set_cell(tbl.cell(7, 2), "無し")

    # Row 8: 研究論文・著書
    pubs = person.get("publications", [])
    set_cell(tbl.cell(8, 2),
             "\n".join(
                 f"{p.get('authors', '')} ({p.get('year', '')}). "
                 f"{p.get('title', '')}. {p.get('journal', '')}."
                 for p in pubs) if pubs else "")

    # Row 9: 知的財産権
    patents = person.get("patents", [])
    if patents:
        set_cell(tbl.cell(9, 2),
                 "\n".join(
                     f"{p.get('title', '')} "
                     f"({p.get('application_number', '')}, {p.get('year', '')})"
                     for p in patents))
    else:
        set_cell(tbl.cell(9, 2), "")

    print(f"  ✓ {label}")


# ============================================================================
# Fill: 参考様式 (consent forms — paragraph-based)
# ============================================================================

def fill_consent_forms(doc, cfg, res):
    """Fill placeholders in 参考様式 (承諾書) sections.

    M15-02: 5 系統の placeholder を YAML から自動 fill する:
      - 「□□ □□」 signer 行（代表機関版のみ）
      - 「研究代表者 所属氏名：」 PI 名（代表機関版のみ）
      - 「研究分担者 所属氏名：」 co_investigator 一覧
      - 「研究課題名：」 課題名
      - 「代表研究機関名：」 代表機関名（分担機関版）
      - 「令和　年度安全保障…」「令和　年度～令和　年度」 元号年置換

    分担研究機関の signer は機関ごとに異なるため自動 fill は対象外
    （手動入力で対応）。日付（令和　年　月　日）も提出時の判断のため未対応。
    """
    inst = cfg["lead_institution"]
    proj = cfg["project"]
    pi = res["pi"]
    co_list = res.get("co_investigators", []) or []

    signer = inst.get("authorized_signer", {}) or {}
    signer_name = signer.get("name", "")
    signer_title = signer.get("title", "")
    inst_name = inst.get("name", "")
    title_ja = proj.get("title_ja", "")

    period_end = str(proj.get("period_end", "R10"))
    # "R10" / "Ｒ10" / "10" / "令和10年度" → "10"
    period_end_year = "".join(c for c in period_end if c.isdigit()) or "10"

    pi_display = (
        f"{pi.get('affiliation', '')}\u3000{pi.get('name_ja', '')}"
    ).strip()
    co_display = "、".join(
        f"{co.get('affiliation', '')}\u3000{co.get('name_ja', '')}".strip()
        for co in co_list
    )
    # N16-07: signer block \u306e placeholder \u69cb\u9020\u306f 3 \u6bb5\uff08\u6a5f\u95a2\u540d / \u8077\u540d / \u6c0f\u540d\uff09\u3002
    # \u25cb\u25cb\u5927\u5b66 / \u25b3\u25b3\u5b66\u90e8\u9577 / \u25a1\u25a1 \u25a1\u25a1 \u3092\u305d\u308c\u305e\u308c\u5225\u6bb5\u843d\u3067\u500b\u5225\u7f6e\u63db\u3059\u308b\u3002
    # 1 \u6bb5\u843d\u306b\u8a70\u3081\u308b\u3068\u30bb\u30eb\u5e45\u3067\u6539\u884c\u3055\u308c\u3066\u53ef\u8aad\u6027\u304c\u52a3\u5316\u3059\u308b\u3002
    lead_inst_display = inst_name
    lead_title_display = signer_title
    lead_signer_display = signer_name

    def _set_para_text(para, new_text):
        """Replace paragraph text while preserving the first run's rPr."""
        rpr = None
        for r in para.runs:
            rpr_el = r._element.find(qn("w:rPr"))
            if rpr_el is not None:
                rpr = copy.deepcopy(rpr_el)
                break
        for r_el in list(para._element.findall(qn("w:r"))):
            para._element.remove(r_el)
        new_r = OxmlElement("w:r")
        if rpr is not None:
            new_r.append(rpr)
        new_t = OxmlElement("w:t")
        new_t.text = new_text
        new_t.set(qn("xml:space"), "preserve")
        new_r.append(new_t)
        para._element.append(new_r)

    def _patch_runs(para):
        """Paragraph-level substitutions (intro 元号年 / 実施期間).

        N16-08: テンプレートは「令和」「\u3000年度」「安全保障\u2026」を
        複数 w:r に分割するため run 単位の str.replace では一致しない。
        paragraph 全体のテキストで判定・置換し、差分があれば
        _set_para_text で 1 ラン化して書き戻す。
        """
        text = para.text
        new = text
        new = new.replace(
            "令和\u3000年度安全保障技術研究推進制度",
            "令和８年度安全保障技術研究推進制度",
        )
        new = new.replace(
            "令和\u3000年度～令和\u3000年度",
            f"令和８年度～令和{period_end_year}年度",
        )
        new = new.replace(
            "令和\u3000年度〜令和\u3000年度",
            f"令和８年度〜令和{period_end_year}年度",
        )
        if new != text:
            _set_para_text(para, new)

    consent_type = None  # "lead" | "partner" | "hojokin" | None
    in_consent = False
    fills = 0

    for para in doc.paragraphs:
        text = para.text
        stripped = text.strip()

        # --- Section boundary ---
        if "（参考様式" in stripped:
            if "代表研究機関" in stripped:
                consent_type = "lead"
            elif "分担研究機関" in stripped:
                consent_type = "partner"
            elif "補助金" in stripped:
                consent_type = "hojokin"
            in_consent = True
            continue
        if "応募・実施承諾書" in stripped:
            if consent_type is None:
                consent_type = "lead"  # 最初の承諾書（heading 不在時）
            in_consent = True
            continue
        if stripped.startswith("（様式") or "チェックリスト" in stripped:
            in_consent = False
            consent_type = None
            continue
        if not in_consent:
            continue

        # --- Paragraph-level fills（lead / partner で挙動を分岐） ---
        # N16-07: signer block 3 段（○○大学 / △△学部長 / □□ □□）を個別置換
        if stripped.startswith("○○") and consent_type == "lead":
            if lead_inst_display:
                _set_para_text(para, lead_inst_display)
                fills += 1
                continue
        if stripped.startswith("△△") and consent_type == "lead":
            if lead_title_display:
                _set_para_text(para, lead_title_display)
                fills += 1
                continue
        if stripped.startswith("□□") and consent_type == "lead":
            if lead_signer_display:
                _set_para_text(para, lead_signer_display)
                fills += 1
                continue
        if (stripped.startswith("研究代表者") and "所属氏名" in stripped
                and pi_display):
            _set_para_text(
                para, f"研究代表者\u3000所属氏名：\u3000{pi_display}"
            )
            fills += 1
            continue
        if (stripped.startswith("研究分担者") and "所属氏名" in stripped
                and co_display):
            _set_para_text(
                para, f"研究分担者\u3000所属氏名：\u3000{co_display}"
            )
            fills += 1
            continue
        if stripped.startswith("研究課題名：") and title_ja:
            _set_para_text(para, f"研究課題名：\u3000{title_ja}")
            fills += 1
            continue
        if stripped.startswith("代表研究機関名：") and inst_name:
            _set_para_text(para, f"代表研究機関名：{inst_name}")
            fills += 1
            continue

        # --- Run-level fills（元号年 placeholder） ---
        _patch_runs(para)

    print(f"  ✓ 参考様式 placeholder fill ({fills} 段落)")


# ============================================================================
# Section deletion
# ============================================================================

def delete_sections(doc, tables, cfg, res):
    """Remove unnecessary forms/sections from the document."""
    proj_type = str(cfg["project"].get("type", "A"))
    inst_type = cfg["lead_institution"].get("type", "")
    has_co = bool(res.get("co_investigators", []))

    body = doc.element.body
    children = list(body)

    def _text(el):
        """Extract text from a paragraph element."""
        if el.tag == qn("w:p"):
            return "".join(t.text or ""
                           for t in el.findall(f".//{qn('w:t')}"))
        return ""

    def _find(text, start=0):
        """Find index of first body-child whose text contains *text*."""
        for i in range(start, len(children)):
            if text in _text(children[i]):
                return i
        return -1

    def _tbl_idx(table):
        """Find body-child index for a python-docx Table."""
        el = table._element
        for i, ch in enumerate(children):
            if ch is el:
                return i
        return -1

    remove = set()

    # --- 1. チェックリスト: always delete ---
    ci = _find("応募書類チェックリスト")
    if ci < 0:
        ci = _find("チェックリスト")
    if ci >= 0:
        for i in range(ci, len(children)):
            if children[i].tag != qn("w:sectPr"):
                remove.add(i)

    # --- 2. 補助金参考様式: delete for 委託事業 (S/A/C) ---
    # M15-05: 「補助事業」「補助金」の単純部分一致は本文中の同単語と誤一致し
    # 削除範囲が肥大する致命的リスク（最悪、様式1-1 以降の全本文を巻き込む）。
    # ここでは「（参考様式」かつ「補助金」を両方含む見出し段落のみを起点とし、
    # 削除範囲は次の「（参考様式」見出し or チェックリスト見出し or w:sectPr
    # 直前 のいずれか早い位置までに限定する。
    if proj_type in ("S", "A", "C"):
        hi = -1
        for i in range(len(children)):
            txt = _text(children[i])
            if "（参考様式" in txt and "補助金" in txt:
                hi = i
                break
        if hi >= 0:
            # 削除終端: 「（参考様式」を含む次の見出し or チェックリスト位置 or
            # w:sectPr 直前 のうち最も早いもの
            end = len(children)
            for i in range(hi + 1, len(children)):
                if children[i].tag == qn("w:sectPr"):
                    end = i
                    break
                txt = _text(children[i])
                if "（参考様式" in txt or "チェックリスト" in txt:
                    end = i
                    break
            if ci >= 0 and ci > hi and ci < end:
                end = ci
            for i in range(hi, end):
                if children[i].tag != qn("w:sectPr"):
                    remove.add(i)

    # --- 2.5. 参考様式 委託費（分担研究機関）: delete if no sub-institutions ---
    # N16-05: sub_institutions=[] の場合、分担研究機関版承諾書は記入欄が
    # 空のまま残り提出書類として誤解を招く（共同研究者指摘）。代表機関版は
    # 残し、分担機関版のみ削除する。区間検出は補助金版と同種の境界条件で限定。
    if not cfg.get("sub_institutions", []):
        hi_p = -1
        for i in range(len(children)):
            txt = _text(children[i])
            if "（参考様式" in txt and "分担研究機関" in txt:
                hi_p = i
                break
        if hi_p >= 0:
            end = len(children)
            for i in range(hi_p + 1, len(children)):
                if children[i].tag == qn("w:sectPr"):
                    end = i
                    break
                txt = _text(children[i])
                if "（参考様式" in txt or "チェックリスト" in txt:
                    end = i
                    break
            if ci >= 0 and ci > hi_p and ci < end:
                end = ci
            for i in range(hi_p, end):
                if children[i].tag != qn("w:sectPr"):
                    remove.add(i)

    # --- 3. 様式5: delete if university/public institution ---
    if inst_type in ("大学等", "公的研究機関"):
        si = _find("（様式５）")
        if si >= 0:
            # M15-05 同種: 削除終端を最も早い境界に揃える
            end = len(children)
            for i in range(si + 1, len(children)):
                if children[i].tag == qn("w:sectPr"):
                    end = i
                    break
                txt = _text(children[i])
                if "（参考様式" in txt or "承諾書" in txt or "（様式" in txt:
                    end = i
                    break
            for i in range(si, end):
                if children[i].tag != qn("w:sectPr"):
                    remove.add(i)

    # --- 4. 様式4-2: delete if no co-investigators ---
    if not has_co:
        fi = _find("（様式４－２）")
        if fi >= 0:
            end = _find("（様式５）", fi + 1)
            if end < 0:
                end = _find("承諾書", fi + 1)
            if end < 0:
                end = fi + 1
            for i in range(fi, end):
                if children[i].tag != qn("w:sectPr"):
                    remove.add(i)

    # --- 5. 様式2-2 years 4 & 5: delete for Type A (3-year) ---
    if proj_type == "A":
        for key in ("2-2_y4", "2-2_y5"):
            if key in tables:
                ti = _tbl_idx(tables[key])
                if ti >= 0:
                    remove.add(ti)
                    # Remove year-header paragraph above the table
                    # (may not be directly above — scan upward past empty paragraphs)
                    for j in range(ti - 1, max(ti - 6, -1), -1):
                        txt = _text(children[j])
                        if "年目" in txt:
                            remove.add(j)
                            break
                        if txt.strip():
                            break  # hit non-empty, non-header content

    # Delete in reverse order to keep indices stable
    for i in sorted(remove, reverse=True):
        if i < len(children) and children[i].tag != qn("w:sectPr"):
            body.remove(children[i])

    if remove:
        print(f"  ✓ Deleted {len(remove)} elements from unnecessary sections")


# ============================================================================
# Main
# ============================================================================

# ============================================================================
# Application date placeholder fill (様式1-1 / 参考様式)
# ============================================================================

def _format_reiwa(d):
    """Convert date(2026,5,10) → '令和8年5月10日'."""
    reiwa_year = d.year - 2018  # 令和元年 = 2019
    return f"令和{reiwa_year}年{d.month}月{d.day}日"


def _parse_submission_date(value):
    """Accept ISO string ('2026-05-10') or datetime.date and return date."""
    if isinstance(value, _date):
        return value
    if isinstance(value, str):
        y, m, d = value.split("-")
        return _date(int(y), int(m), int(d))
    raise ValueError(f"submission_date must be ISO string or date, got {value!r}")


_DATE_PLACEHOLDER_RE = re.compile(r"令和[　\s]*年[　\s]*月[　\s]*日")


def fill_application_dates(doc, cfg):
    """Replace '令和　年　月　日' placeholders with submission_date.

    Targets standalone date paragraphs in 様式1-1 直下 and 参考様式
    （応募・実施承諾書）末尾。本文中の「令和　年度〜令和　年度」のような
    元号年プレースホルダは対象外（fill_consent_forms 側で処理済）。
    """
    proj = cfg.get("project", {})
    raw = proj.get("submission_date")
    if not raw:
        return
    try:
        d = _parse_submission_date(raw)
    except ValueError as e:
        warnings.warn(f"submission_date parse failed: {e}")
        return
    formatted = _format_reiwa(d)

    n = 0
    for para in doc.paragraphs:
        text = para.text
        if not _DATE_PLACEHOLDER_RE.search(text):
            continue
        new_text = _DATE_PLACEHOLDER_RE.sub(formatted, text)
        if new_text == text:
            continue
        # Preserve first run's rPr to retain font/size formatting
        rpr = None
        for r in para.runs:
            rpr_el = r._element.find(qn("w:rPr"))
            if rpr_el is not None:
                rpr = copy.deepcopy(rpr_el)
                break
        for r_el in list(para._element.findall(qn("w:r"))):
            para._element.remove(r_el)
        new_r = OxmlElement("w:r")
        if rpr is not None:
            new_r.append(rpr)
        new_t = OxmlElement("w:t")
        new_t.text = new_text
        new_t.set(qn("xml:space"), "preserve")
        new_r.append(new_t)
        para._element.append(new_r)
        n += 1
    print(f"  ✓ 申請日 placeholder fill ({n} 段落, {formatted})")


# ============================================================================
# Title / 研究者氏名 placeholder fill (様式1-2 / 1-3 / 3-1 / 3-2)
# ============================================================================

def _replace_para_text(para, new_text):
    """Replace paragraph text while preserving the first run's rPr."""
    rpr = None
    for r in para.runs:
        rpr_el = r._element.find(qn("w:rPr"))
        if rpr_el is not None:
            rpr = copy.deepcopy(rpr_el)
            break
    for r_el in list(para._element.findall(qn("w:r"))):
        para._element.remove(r_el)
    new_r = OxmlElement("w:r")
    if rpr is not None:
        new_r.append(rpr)
    new_t = OxmlElement("w:t")
    new_t.text = new_text
    new_t.set(qn("xml:space"), "preserve")
    new_r.append(new_t)
    para._element.append(new_r)


def fill_title_placeholders(doc, cfg):
    """Fill standalone '研究課題名：' placeholder paragraphs with project title.

    fill_consent_forms 側は参考様式区間内のみを対象とするため、様式1-2 /
    様式1-3 直下の同じプレースホルダ段落は未処理のまま残る。本関数で全文書を
    走査し、空欄状態の研究課題名行を全て課題名で埋める。既に内容が入っている
    段落（例: 参考様式で fill 済み）は idempotent に skip する。
    """
    proj = cfg.get("project", {})
    title_ja = proj.get("title_ja", "")
    if not title_ja:
        return
    target = f"研究課題名：　{title_ja}"
    n = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text.startswith("研究課題名："):
            continue
        rest = text[len("研究課題名："):].strip()
        # idempotent skip: 既に課題名が入っている段落はそのまま
        if title_ja in rest:
            continue
        # 空欄プレースホルダ（全角空白のみ等）を fill 対象とする
        _replace_para_text(para, target)
        n += 1
    print(f"  ✓ 研究課題名 placeholder fill ({n} 段落)")


def fill_form3_researcher_names(doc, cfg, res):
    """Fill '研究代表者：' / '研究分担者：' placeholders in 様式3-1 / 3-2."""
    pi = res.get("pi", {})
    pi_name = pi.get("name_ja", "")
    co_list = res.get("co_investigators", []) or []
    co_names = "、".join(co.get("name_ja", "") for co in co_list if co.get("name_ja"))

    n = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("研究代表者：") and pi_name:
            rest = text[len("研究代表者："):].strip()
            if pi_name in rest:
                continue
            _replace_para_text(para, f"研究代表者：　{pi_name}")
            n += 1
        elif text.startswith("研究分担者：") and co_names:
            rest = text[len("研究分担者："):].strip()
            if all(co.get("name_ja", "") in rest for co in co_list if co.get("name_ja")):
                continue
            _replace_para_text(para, f"研究分担者：　{co_names}")
            n += 1
    print(f"  ✓ 様式3 研究者氏名 placeholder fill ({n} 段落)")


# ============================================================================
# Trailing empty paragraph cleanup (blank page removal)
# ============================================================================

def remove_trailing_empty_before_page_breaks(doc):
    """Remove runs of empty paragraphs that immediately precede a form heading
    with pageBreakBefore.

    N16-04: テンプレート（r08youshiki1_5.docx）は様式間に大量の空段落
    （様式2-1 末尾に 16 段落、2-2 末尾に 14 段落 等）を持ち、insert_form_page_breaks
    が次の様式ヘッダに pageBreakBefore を入れる際、この trailing 空段落が
    そのまま新ページに繰り越されて blank page が発生する（共同研究者指摘）。

    各様式ヘッダの直前を逆走し、空 (text のみが空白) かつ table を含まない
    段落を順次削除する。最初に非空段落・テーブル・別の様式ヘッダに到達した
    時点で停止する。
    """
    body = doc.element.body
    children = list(body)

    def _text(el):
        if el.tag != qn("w:p"):
            return ""
        return "".join(t.text or "" for t in el.findall(f".//{qn('w:t')}")).strip()

    def _has_pbb(el):
        if el.tag != qn("w:p"):
            return False
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            return False
        return pPr.find(qn("w:pageBreakBefore")) is not None

    removed = 0
    for el in list(children):
        if not _has_pbb(el):
            continue
        idx = list(body).index(el)
        # walk backwards, removing empty paragraphs
        j = idx - 1
        while j >= 0:
            prev = list(body)[j]
            if prev.tag != qn("w:p"):
                break
            if _text(prev):
                break
            # safety: don't cross another form heading
            body.remove(prev)
            removed += 1
            j -= 1
    print(f"  ✓ trailing 空段落削除 ({removed} 段落)")


def insert_form_page_breaks(doc):
    """Add w:pageBreakBefore to each 様式 / 参考様式 header paragraph.

    Skips the very first header (様式１－１) since it's the document front
    and an extra page break would produce a blank first page. Idempotent:
    if pPr already has pageBreakBefore, leaves it alone.

    Pattern matches:
      - "（様式1-1）" / "（様式１－２）" / "（様式２－２）" etc.
      - "（参考様式 委託費（代表研究機関））" 等
    Full-width paren "（" is required (the template uses full-width).
    """
    import re as _re
    pattern = _re.compile(r"^（\s*(様式|参考様式)")
    matched_n = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not pattern.match(text):
            continue
        matched_n += 1
        if matched_n == 1:
            continue  # 先頭 様式1-1 は skip
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            para._p.insert(0, pPr)
        if pPr.find(qn("w:pageBreakBefore")) is None:
            pPr.append(OxmlElement("w:pageBreakBefore"))
    print(f"  pageBreakBefore を {max(0, matched_n - 1)} 様式ヘッダに注入")


def main():
    ap = argparse.ArgumentParser(
        description="Fill table forms in r08youshiki1_5.docx")
    ap.add_argument("--config", default="main/00_setup/config.yaml",
                    help="config.yaml path")
    ap.add_argument("--researchers", default="main/00_setup/researchers.yaml",
                    help="researchers.yaml path")
    ap.add_argument("--other-funding", default="main/00_setup/other_funding.yaml",
                    help="other_funding.yaml path")
    ap.add_argument("--source", default="data/source/r08youshiki1_5.docx",
                    help="Source docx template")
    ap.add_argument("--output", default="main/step02_docx/output/",
                    help="Output directory")
    args = ap.parse_args()

    # ---- Load data ----
    print("Loading YAML data...")
    cfg = load_yaml(args.config)
    res = load_yaml(args.researchers)
    ofund = load_yaml(args.other_funding)

    # ---- Copy source ----
    out_dir = Path(args.output)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "youshiki1_5_filled.docx"
    shutil.copy2(args.source, out_path)
    print(f"Copied template -> {out_path}")

    # ---- Open ----
    doc = Document(str(out_path))

    # ---- Identify tables ----
    tables = identify_tables(doc)
    print(f"Identified tables: {sorted(tables.keys())}")

    has_co = bool(res.get("co_investigators", []))

    # ---- Fill forms ----
    print("\nFilling forms...")

    if "1-1" in tables:
        fill_1_1(tables["1-1"], cfg, res, ofund)
    else:
        warnings.warn("様式1-1 not found")

    if "2-1" in tables:
        fill_2_1(tables["2-1"], cfg)

    if "2-1_inst" in tables:
        fill_2_1_inst(tables["2-1_inst"], cfg)

    fill_2_2(tables, cfg)

    if "3-1" in tables:
        fill_3_1(tables["3-1"], cfg, res, ofund)

    if "3-2" in tables:
        fill_3_2(tables["3-2"], cfg, res, ofund)

    if "4-1" in tables:
        fill_4(tables["4-1"], cfg, res["pi"], "様式4-1")

    if "4-2" in tables and has_co:
        co_list = res["co_investigators"]
        fill_4(tables["4-2"], cfg, co_list[0],
               f"様式4-2 (1/{len(co_list)} {co_list[0]['name_ja']})")
        # Duplicate 4-2 table for remaining co-investigators
        prev_el = tables["4-2"]._element
        for ci in range(1, len(co_list)):
            # Page break paragraph between tables
            bp = OxmlElement("w:p")
            br_run = OxmlElement("w:r")
            br_el = OxmlElement("w:br")
            br_el.set(qn("w:type"), "page")
            br_run.append(br_el)
            bp.append(br_run)
            prev_el.addnext(bp)
            # Deep-copy and insert table
            new_tbl = copy.deepcopy(tables["4-2"]._element)
            bp.addnext(new_tbl)
            # M15-03: parent は元 4-2 テーブルと同じ python-docx の親
            # （通常 _Body or _Cell）を使う。素の lxml 要素 (doc.element.body) を
            # 渡すと .part プロパティが解決できず、cell 拡張時に AttributeError。
            wrapped = DocxTable(new_tbl, tables["4-2"]._parent)
            fill_4(wrapped, cfg, co_list[ci],
                   f"様式4-2 ({ci + 1}/{len(co_list)} {co_list[ci]['name_ja']})")
            prev_el = new_tbl

    fill_consent_forms(doc, cfg, res)

    # ---- Title / 研究者氏名 placeholder fill (様式1-2 / 1-3 / 3-1 / 3-2) ----
    fill_title_placeholders(doc, cfg)
    fill_form3_researcher_names(doc, cfg, res)

    # ---- Application date fill (様式1-1 / 参考様式) ----
    fill_application_dates(doc, cfg)

    # ---- Delete unnecessary sections ----
    print("\nCleaning up...")
    delete_sections(doc, tables, cfg, res)

    # ---- Enforce page break between forms ----
    # N15-03: テンプレート r08youshiki1_5.docx は 様式間に明示的な
    # pageBreakBefore を持たないため、ナラティブ挿入後のページ流れが
    # 様式途中から次様式に繋がる。各様式ヘッダ（（様式X－Y） / （参考様式...））
    # に pageBreakBefore を注入することで常に新ページから開始させる。
    # 先頭の（様式１－１）は文頭なのでスキップ。
    insert_form_page_breaks(doc)

    # ---- Remove trailing empty paragraphs that cause blank pages ----
    # N16-04: pageBreakBefore 直前の空段落は新ページに繰り越されて blank page
    # を生む（共同研究者指摘）。各 pageBreakBefore 段落の直前を逆走し空段落を
    # 削除する。insert_form_page_breaks の後に必ず実行すること。
    remove_trailing_empty_before_page_breaks(doc)

    # ---- Save ----
    doc.save(str(out_path))
    print(f"\nDone: {out_path}")


if __name__ == "__main__":
    main()
