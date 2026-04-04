#!/usr/bin/env python3
"""fill_security.py — 別紙5・別添 記入スクリプト

data/source/r08youshiki_besshi5.docx → 別紙5（研究セキュリティ質問票）
data/source/r08youshiki_betten.docx  → 別添（自己申告書）× 研究者人数分

Usage:
    python fill_security.py \
        --config main/00_setup/config.yaml \
        --researchers main/00_setup/researchers.yaml \
        --security main/00_setup/security.yaml \
        --besshi5 data/source/r08youshiki_besshi5.docx \
        --betten data/source/r08youshiki_betten.docx \
        --output main/step02_docx/output/
"""

import argparse
import copy
import shutil
import warnings
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# YAML helpers
# ============================================================================

def load_yaml(path):
    """Load a YAML file and return its contents."""
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


# ============================================================================
# Cell-level helpers
# ============================================================================

def _first_rpr(cell):
    """Return deepcopy of the first run's rPr, or None."""
    for p in cell.paragraphs:
        for r in p.runs:
            rpr = r._element.find(qn("w:rPr"))
            if rpr is not None:
                return copy.deepcopy(rpr)
    return None


def _first_ppr(cell):
    """Return deepcopy of the first paragraph's pPr, or None."""
    if cell.paragraphs:
        ppr = cell.paragraphs[0]._element.find(qn("w:pPr"))
        if ppr is not None:
            return copy.deepcopy(ppr)
    return None


def set_cell(cell, text):
    r"""Replace cell content.  ``\n`` → new paragraph.  Preserves first-run format."""
    text = str(text) if text is not None else ""
    rpr = _first_rpr(cell)
    ppr = _first_ppr(cell)
    tc = cell._element
    for old_p in list(tc.findall(qn("w:p"))):
        tc.remove(old_p)
    for line in (text.split("\n") if text else [""]):
        p = OxmlElement("w:p")
        if ppr is not None:
            p.append(copy.deepcopy(ppr))
        r = OxmlElement("w:r")
        if rpr is not None:
            r.append(copy.deepcopy(rpr))
        t = OxmlElement("w:t")
        t.text = line
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)


# ============================================================================
# Checkbox helper
# ============================================================================

def check_box(cell, target):
    """Replace □ with ☑ for the checkbox option containing *target*.

    Works by finding the paragraph that contains the target text,
    locating the nearest □ before it, and replacing in the specific run
    to preserve formatting.
    """
    for p in cell.paragraphs:
        if target not in p.text or "□" not in p.text:
            continue
        p_text = p.text
        target_pos = p_text.index(target)
        box_pos = p_text.rfind("□", 0, target_pos + 1)
        if box_pos == -1:
            continue
        # Find the run containing this □
        offset = 0
        for run in p.runs:
            run_len = len(run.text)
            if offset <= box_pos < offset + run_len:
                local = box_pos - offset
                run.text = run.text[:local] + "☑" + run.text[local + 1:]
                return True
            offset += run_len

    # Fallback: line-level replacement via set_cell
    full = cell.text
    for line in full.split("\n"):
        if target in line and "□" in line:
            new_full = full.replace(line, line.replace("□", "☑", 1), 1)
            set_cell(cell, new_full)
            return True

    warnings.warn(f"check_box: '{target}' not found in cell")
    return False


# ============================================================================
# Row-addition helper
# ============================================================================

def add_table_rows(table, count):
    """Append *count* rows to *table*, deep-copied from the last existing row."""
    if count <= 0:
        return
    template_tr = table.rows[-1]._tr
    for _ in range(count):
        new_tr = copy.deepcopy(template_tr)
        table._tbl.append(new_tr)


# ============================================================================
# 13-item configuration
# ============================================================================

# (yaml_key, has_header_row, col_keys, item_type)
ITEMS = [
    # ① 学歴
    ("education_history",    False, ["year_month", "description"], "list"),
    # ② 職歴
    ("career_history",       False, ["year_month", "description"], "list"),
    # ③ 研究費受給歴
    ("funding_history",      True,  ["date", "program_name", "agency", "period", "total_amount"], "list"),
    # ④ 研究費以外の支援
    ("non_research_support", True,  ["year_month", "description", "agency", "period", "total_amount"], "list"),
    # ⑤ 発表論文
    ("publications",         True,  ["date", "title", "first_author", "corresponding_author", "co_authors"], "list"),
    # ⑥ 特許
    ("patents",              True,  ["date", "description", "applicant", "co_inventors", "co_applicants"], "list"),
    # ⑦ 外国の人材採用プログラム
    ("foreign_talent_programs", True, ["period", "program_name", "host_country"], "list"),
    # ⑧ 処分歴
    ("disciplinary_history", True,  ["year_month", "description"], "list"),
    # ⑨ リスト掲載有無
    ("list_status",          True,  None, "status"),
    # ⑩ リスト掲載機関への所属歴
    ("listed_entity_affiliation", True, ["period", "description"], "list"),
    # ⑪ リスト掲載機関の研究者との関係
    ("listed_entity_relationships", True, ["year_month", "description", "entity_name"], "list"),
    # ⑫ 非居住者・特定類型
    ("residency_status",     True,  None, "status"),
    # ⑬ 国籍
    ("nationality",          False, None, "nationality"),
]

# Due-diligence status YAML value → checkbox label
DD_STATUS_MAP = {
    "提出済":       "提出した",
    "要提出・未提出": "提出は必要であるが提出していない",
    "不要":         "提出は不要である",
}

# Due-diligence result YAML value → checkbox label
DD_RESULT_MAP = {
    "実施した":          "実施した",
    "実施していない":      "実施していない",
    "一部取得できなかった": "実施したが一部の情報は取得できなかった",
}

# Risk-measure type → substring to find in Table 21 checkboxes
MEASURE_LABEL = {
    "facility_access":      "施設・設備へのアクセス権限の管理",
    "research_location":    "オフキャンパス",
    "meeting_participants": "ミーティング等への参加者",
    "employment_contract":  "雇用契約を締結する",
    "training":             "研修の受講",
    "data_access":          "研究データ等の情報へのアクセス権限",
    "cybersecurity":        "サイバー攻撃",
    "other":                "その他",
}


# ============================================================================
# Researcher list builder
# ============================================================================

def build_researcher_list(res_yaml):
    """Return [(name_ja, name_en, affiliation, department, position), ...]."""
    pi = res_yaml["pi"]
    out = [(
        pi["name_ja"], pi["name_en"],
        pi.get("affiliation", ""), pi.get("department", ""), pi.get("position", ""),
    )]
    for co in res_yaml.get("co_investigators", []):
        out.append((
            co["name_ja"], co["name_en"],
            co.get("affiliation", ""), co.get("department", ""), co.get("position", ""),
        ))
    return out


# ============================================================================
# Fill list-type item into a table
# ============================================================================

def _fill_list_item(table, entries, col_keys, data_start):
    """Fill *entries* into *table* from row *data_start*, adding rows if needed."""
    if not entries:
        return
    avail = len(table.rows) - data_start
    if len(entries) > avail:
        add_table_rows(table, len(entries) - avail)
    ncols = len(col_keys)
    for ri, entry in enumerate(entries):
        row = data_start + ri
        for ci in range(ncols):
            val = entry.get(col_keys[ci], "") if isinstance(entry, dict) else ""
            set_cell(table.cell(row, ci), str(val))
    # Clear remaining placeholder rows
    for ri in range(data_start + len(entries), len(table.rows)):
        for ci in range(len(table.columns)):
            set_cell(table.cell(ri, ci), "")


# ============================================================================
# Fill: 別紙5  (24 tables)
# ============================================================================

def fill_besshi5(doc, cfg, res_yaml, sec):
    """Fill r08youshiki_besshi5.docx."""
    tables = doc.tables
    pi = res_yaml["pi"]
    proj = cfg["project"]
    researchers = build_researcher_list(res_yaml)
    sec_r = sec.get("researchers", {})
    dd = sec.get("due_diligence", {})
    risk = sec.get("risk_assessment", {})
    consent_data = sec.get("consent", {})

    # --- Table 0: 提案情報 (3r×2c) ---
    set_cell(tables[0].cell(0, 1), pi["name_ja"])
    set_cell(tables[0].cell(1, 1),
             f"{pi.get('affiliation', '')}・"
             f"{pi.get('department', '')}・"
             f"{pi.get('position', '')}")
    set_cell(tables[0].cell(2, 1), proj["title_ja"])
    print("  ✓ Table 0: 提案情報")

    # --- Tables 1–2: §1(1) デューデリジェンス状況 ---
    lead = dd.get("lead_institution", {})
    lead_label = DD_STATUS_MAP.get(lead.get("status", ""), "")
    if lead_label:
        check_box(tables[1].cell(0, 0), lead_label)

    partners = dd.get("partner_institutions", [])
    if partners:
        p_label = DD_STATUS_MAP.get(partners[0].get("status", ""), "")
        if p_label:
            check_box(tables[2].cell(0, 0), p_label)
    print("  ✓ Tables 1-2: デューデリジェンス状況")

    # --- Tables 3–15: §1(2) 13項目 ---
    for item_idx, (key, has_hdr, col_keys, itype) in enumerate(ITEMS):
        tbl = tables[3 + item_idx]
        ds = 1 if has_hdr else 0  # data-start row

        if itype == "nationality":
            # Table 15 (1r×2c): list all researchers' nationalities
            parts = []
            for name_ja, *_ in researchers:
                nat = sec_r.get(name_ja, {}).get("nationality", "")
                parts.append(f"{name_ja}: {nat}")
            set_cell(tbl.cell(0, 1), "\n".join(parts))

        elif itype == "status":
            # ⑨ list_status / ⑫ residency_status: one row per researcher
            rows = []
            for name_ja, *_ in researchers:
                sd = sec_r.get(name_ja, {}).get(key, {})
                appl = sd.get("applicable", False)
                det = sd.get("details", "")
                appl_str = "有" if appl else "無"
                det_str = f"（{name_ja}）{det}" if appl else f"（{name_ja}）"
                rows.append((appl_str, det_str))
            avail = len(tbl.rows) - ds
            if len(rows) > avail:
                add_table_rows(tbl, len(rows) - avail)
            for ri, (a, d) in enumerate(rows):
                set_cell(tbl.cell(ds + ri, 0), a)
                set_cell(tbl.cell(ds + ri, 1), d)
            for ri in range(ds + len(rows), len(tbl.rows)):
                for ci in range(len(tbl.columns)):
                    set_cell(tbl.cell(ri, ci), "")

        else:  # "list"
            all_entries = []
            for name_ja, *_ in researchers:
                data = sec_r.get(name_ja, {}).get(key, [])
                if isinstance(data, list) and data:
                    for ei, entry in enumerate(data):
                        tagged = dict(entry)
                        if ei == 0:
                            first_key = col_keys[0]
                            tagged[first_key] = (
                                f"【{name_ja}】\n{tagged.get(first_key, '')}")
                        all_entries.append(tagged)
                else:
                    # No data → "該当なし" marker
                    all_entries.append({col_keys[0]: f"【{name_ja}】該当なし"})

            # If every researcher has no data, leave table untouched
            if all("該当なし" in e.get(col_keys[0], "") for e in all_entries):
                pass
            else:
                _fill_list_item(tbl, all_entries, col_keys, ds)

    print("  ✓ Tables 3-15: 13項目")

    # --- Table 16: §1 デューデリジェンス結果 ---
    result = dd.get("result", "")
    result_label = DD_RESULT_MAP.get(result, result)
    if result_label:
        check_box(tables[16].cell(0, 0), result_label)
    print("  ✓ Table 16: デューデリジェンス結果")

    # --- Table 17: §2(2) リスク軽減措置要否 ---
    check_box(tables[17].cell(0, 0),
              "はい" if risk.get("mitigation_needed") else "いいえ")
    print("  ✓ Table 17: リスク軽減措置要否")

    # --- Table 18: §2(3) 新規追加時確認 ---
    if risk.get("new_member_confirmed"):
        check_box(tables[18].cell(0, 0), "記載内容に従います")
    print("  ✓ Table 18: 新規追加時確認")

    # --- Tables 19–20: §3 共同研究機関リスク ---
    pr = risk.get("partner_institution_risk", {})
    check_box(tables[19].cell(0, 0),
              "はい" if pr.get("has_risk") else "いいえ")
    check_box(tables[20].cell(0, 0),
              "はい" if pr.get("has_risk_researcher") else "いいえ")
    print("  ✓ Tables 19-20: 共同研究機関リスク")

    # --- Table 21: §4 リスク軽減措置 ---
    for m in risk.get("measures", []):
        label = MEASURE_LABEL.get(m.get("type", ""), "")
        if label:
            check_box(tables[21].cell(0, 0), label)
    print("  ✓ Table 21: リスク軽減措置")

    # --- Table 22: §4(2) 実行確認 ---
    if risk.get("measures_confirmed"):
        check_box(tables[22].cell(0, 0), "実行可能であることを確認しました")
    print("  ✓ Table 22: 実行確認")

    # --- Table 23: 同意確認 ---
    if consent_data.get("all_members_consented"):
        check_box(tables[23].cell(0, 0),
                  "研究体制におけるPI、Co-PI及び研究参画者の同意を得ました")
    if consent_data.get("institutional_confirmation"):
        check_box(tables[23].cell(0, 0),
                  "研究代表機関及び共同研究機関の担当部署の確認を得ました")
    print("  ✓ Table 23: 同意確認")


# ============================================================================
# Fill: 別添  (per researcher)
# ============================================================================

def fill_betten(doc, sec, name_ja, affiliation, department, position):
    """Fill betten (self-declaration) for one researcher."""
    tables = doc.tables
    r_sec = sec.get("researchers", {}).get(name_ja, {})

    # --- Header: 氏名・所属 ---
    for p in doc.paragraphs:
        if "氏" in p.text and "名" in p.text:
            rpr = None
            for r in p.runs:
                rpr_el = r._element.find(qn("w:rPr"))
                if rpr_el is not None:
                    rpr = copy.deepcopy(rpr_el)
                    break
            for r_el in list(p._element.findall(qn("w:r"))):
                p._element.remove(r_el)
            new_r = OxmlElement("w:r")
            if rpr is not None:
                new_r.append(rpr)
            new_t = OxmlElement("w:t")
            new_t.text = (f"氏　名　{name_ja}"
                          f"（{affiliation}・{department}・{position}）")
            new_t.set(qn("xml:space"), "preserve")
            new_r.append(new_t)
            p._element.append(new_r)
            break

    # --- Tables 0–12: 13項目 ---
    for item_idx, (key, has_hdr, col_keys, itype) in enumerate(ITEMS):
        tbl = tables[item_idx]
        ds = 1 if has_hdr else 0

        if itype == "nationality":
            set_cell(tbl.cell(0, 1), r_sec.get("nationality", ""))

        elif itype == "status":
            sd = r_sec.get(key, {})
            appl = sd.get("applicable", False)
            det = sd.get("details", "")
            set_cell(tbl.cell(ds, 0), "有" if appl else "無")
            set_cell(tbl.cell(ds, 1), det if appl else "")
            for ri in range(ds + 1, len(tbl.rows)):
                for ci in range(len(tbl.columns)):
                    set_cell(tbl.cell(ri, ci), "")

        else:  # "list"
            entries = r_sec.get(key, [])
            if isinstance(entries, list) and entries:
                _fill_list_item(tbl, entries, col_keys, ds)
            # else: leave table as-is (「該当なし」)

    print(f"  ✓ 別添: {name_ja}")


# ============================================================================
# Romanized family name extraction
# ============================================================================

def _family_name(name_en):
    """Extract lowercase romanized family name from 'Given FAMILY' format."""
    for part in name_en.split():
        if part.isupper() and len(part) > 1:
            return part.lower()
    parts = name_en.split()
    return parts[-1].lower() if parts else "unknown"


# ============================================================================
# Main
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Fill 別紙5 (besshi5) and 別添 (betten) security forms")
    parser.add_argument("--config", default="main/00_setup/config.yaml")
    parser.add_argument("--researchers", default="main/00_setup/researchers.yaml")
    parser.add_argument("--security", default="main/00_setup/security.yaml")
    parser.add_argument("--besshi5", default="data/source/r08youshiki_besshi5.docx")
    parser.add_argument("--betten", default="data/source/r08youshiki_betten.docx")
    parser.add_argument("--output", default="main/step02_docx/output/")
    args = parser.parse_args()

    cfg = load_yaml(args.config)
    res = load_yaml(args.researchers)
    sec = load_yaml(args.security)

    outdir = Path(args.output)
    outdir.mkdir(parents=True, exist_ok=True)
    researchers = build_researcher_list(res)

    # === 別紙5 ===
    print("=== 別紙5: 研究セキュリティ質問票 ===")
    besshi5_out = outdir / "besshi5_filled.docx"
    shutil.copy2(args.besshi5, besshi5_out)
    doc5 = Document(str(besshi5_out))
    fill_besshi5(doc5, cfg, res, sec)
    doc5.save(str(besshi5_out))
    print(f"  → {besshi5_out}")

    # === 別添 ===
    print("\n=== 別添: セキュリティ自己申告書 ===")
    for idx, (name_ja, name_en, affil, dept, pos) in enumerate(researchers):
        family = _family_name(name_en)
        fname = f"betten_{idx + 1:02d}_{family}.docx"
        betten_out = outdir / fname
        shutil.copy2(args.betten, betten_out)
        doc_b = Document(str(betten_out))
        fill_betten(doc_b, sec, name_ja, affil, dept, pos)
        doc_b.save(str(betten_out))
        print(f"  → {betten_out}")

    print("\n✓ 完了")


if __name__ == "__main__":
    main()
