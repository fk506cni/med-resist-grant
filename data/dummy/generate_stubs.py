#!/usr/bin/env python3
"""generate_stubs.py — E2Eテスト用スタブ docx/xlsx を data/dummy/ に生成する

fill_forms.py / fill_security.py / fill_excel.py が期待するテーブル構造を
最小限で再現する。
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt
import openpyxl


DUMMY_DIR = Path(__file__).resolve().parent


# ============================================================================
# Helper
# ============================================================================

def _add_table(doc, rows, cols, first_cell_text=""):
    """Add a table to doc and set cell(0,0) text."""
    tbl = doc.add_table(rows=rows, cols=cols)
    if first_cell_text:
        tbl.cell(0, 0).text = first_cell_text
    return tbl


# ============================================================================
# r08youshiki1_5.docx — 様式1-1〜5 + 参考様式 + チェックリスト
#
# identify_tables() expects:
#   1-1:      20r×8c  cell(0,0) contains "①研究テーマ"
#   2-1:       9r×7c  some row col-0 contains "ア．物品費"
#   2-1_inst:  6r×7c  cell(0,0) contains "研究費の内訳", after 2-1
#   2-2_y1..3: 12r×5c cell(1,0) contains "Ⅰ" or "物品費"  (×3 years)
#   3-1:       6r×8c  some row(0..2) col-2 contains "【本研究課題】"
#   3-2:       6r×8c  same pattern (second match)
#   4-1:      10r×5c  cell(0,0) contains "研究課題名"
#   4-2:      10r×5c  same pattern (second match)
#   5:        26r×7c  cell(0,0) starts with "企"
#   checklist: Nx2   header row has "チェック"
#
# Also needs paragraphs with "承諾書" for consent form processing.
# ============================================================================

def create_youshiki1_5():
    doc = Document()

    # -- 様式1-1: 20r×8c --
    tbl = _add_table(doc, 20, 8, "①研究テーマ")
    # Row 0-19 exist; fill some structural cells
    # Row 17: ⑪ 研究分担者 (for fill_1_1)
    tbl.cell(17, 0).text = "⑪研究分担者"

    doc.add_paragraph("")  # separator

    # -- 様式1-2 セクション（inject_narrative.py が検出するマーカー） --
    doc.add_paragraph("（様式１－２）")
    doc.add_paragraph("１．研究目的")
    doc.add_paragraph("ここに記載してください")

    # -- 様式1-3 セクション --
    doc.add_paragraph("（様式１－３）")
    doc.add_paragraph("（１）研究概要")
    doc.add_paragraph("ここに記載してください")

    # -- 様式2-1 セクション --
    doc.add_paragraph("（様式２－１）")

    # -- 様式2-1 費目内訳: 9r×7c --
    tbl = _add_table(doc, 9, 7)
    tbl.cell(1, 0).text = "ア．物品費"  # identify as 2-1

    doc.add_paragraph("")

    # -- 様式2-1 機関別: 6r×7c --
    tbl = _add_table(doc, 6, 7, "研究費の内訳")

    doc.add_paragraph("")

    # -- 様式2-2 ×3 years: 12r×5c each --
    for i in range(3):
        tbl = _add_table(doc, 12, 5)
        tbl.cell(0, 0).text = f"年度 {i+1}"
        tbl.cell(1, 0).text = "Ⅰ．物品費"

        doc.add_paragraph("")

    # -- 様式3-1: 6r×8c --
    tbl = _add_table(doc, 6, 8)
    tbl.cell(1, 2).text = "【本研究課題】"

    doc.add_paragraph("")

    # -- 様式3-2: 6r×8c --
    tbl = _add_table(doc, 6, 8)
    tbl.cell(1, 2).text = "【本研究課題】"

    doc.add_paragraph("")

    # -- 様式4-1: 10r×5c --
    tbl = _add_table(doc, 10, 5, "研究課題名")

    doc.add_paragraph("")

    # -- 様式4-2: 10r×5c --
    tbl = _add_table(doc, 10, 5, "研究課題名")

    doc.add_paragraph("")

    # -- 様式5: 26r×7c --
    tbl = _add_table(doc, 26, 7, "企業名")

    doc.add_paragraph("")

    # -- 参考様式（承諾書） paragraphs --
    doc.add_paragraph("（参考様式）")
    doc.add_paragraph("承諾書（委託契約・代表研究機関）")
    doc.add_paragraph("R8～R\u3000年度")
    doc.add_paragraph("")
    doc.add_paragraph("承諾書（委託契約・分担研究機関）")
    doc.add_paragraph("R8～R\u3000年度")
    doc.add_paragraph("")
    # 補助金版（Type A では削除対象）
    doc.add_paragraph("（参考様式）")
    doc.add_paragraph("承諾書（補助金）")
    doc.add_paragraph("R8～R\u3000年度")

    doc.add_paragraph("")

    # -- チェックリスト: Nx2 --
    tbl = _add_table(doc, 10, 2)
    tbl.cell(0, 0).text = "チェック項目"
    tbl.cell(0, 1).text = "チェック"

    out = DUMMY_DIR / "r08youshiki1_5.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")


# ============================================================================
# r08youshiki_besshi5.docx — 別紙5: 24 tables
# ============================================================================

def create_besshi5():
    doc = Document()

    # Table 0: 提案情報 (3r×2c)
    tbl = _add_table(doc, 3, 2, "研究代表者名")
    tbl.cell(1, 0).text = "所属"
    tbl.cell(2, 0).text = "研究課題名"

    # Tables 1-2: §1(1) デューデリジェンス (each 1r×1c with □ checkboxes)
    for i in range(2):
        tbl = _add_table(doc, 1, 1)
        tbl.cell(0, 0).text = (
            "□ 提出した\n"
            "□ 提出は必要であるが提出していない\n"
            "□ 提出は不要である"
        )

    # Tables 3-15: §1(2) 13項目
    # Match ITEMS structure from fill_security.py
    table_specs = [
        # (rows, cols, has_header)
        (3, 2, False),   # ① education_history
        (3, 2, False),   # ② career_history
        (3, 5, True),    # ③ funding_history
        (3, 5, True),    # ④ non_research_support
        (3, 5, True),    # ⑤ publications
        (3, 5, True),    # ⑥ patents
        (3, 3, True),    # ⑦ foreign_talent_programs
        (3, 2, True),    # ⑧ disciplinary_history
        (3, 2, True),    # ⑨ list_status
        (3, 2, True),    # ⑩ listed_entity_affiliation
        (3, 3, True),    # ⑪ listed_entity_relationships
        (3, 2, True),    # ⑫ residency_status
        (1, 2, False),   # ⑬ nationality
    ]

    for rows, cols, has_header in table_specs:
        total_rows = rows + (1 if has_header else 0)
        tbl = _add_table(doc, total_rows, cols)
        if has_header:
            tbl.cell(0, 0).text = "ヘッダー"

    # Table 16: §1 デューデリジェンス結果 (1r×1c)
    tbl = _add_table(doc, 1, 1)
    tbl.cell(0, 0).text = (
        "□ 実施した\n"
        "□ 実施していない\n"
        "□ 実施したが一部の情報は取得できなかった"
    )

    # Table 17: §2(2) リスク軽減措置要否 (1r×1c)
    tbl = _add_table(doc, 1, 1)
    tbl.cell(0, 0).text = "□ はい　□ いいえ"

    # Table 18: §2(3) 新規追加時確認 (1r×1c)
    tbl = _add_table(doc, 1, 1)
    tbl.cell(0, 0).text = "□ 記載内容に従います"

    # Tables 19-20: §3 共同研究機関リスク (each 1r×1c)
    for _ in range(2):
        tbl = _add_table(doc, 1, 1)
        tbl.cell(0, 0).text = "□ はい　□ いいえ"

    # Table 21: §4 リスク軽減措置 (1r×1c)
    tbl = _add_table(doc, 1, 1)
    tbl.cell(0, 0).text = (
        "□ 施設・設備へのアクセス権限の管理\n"
        "□ オフキャンパス\n"
        "□ ミーティング等への参加者\n"
        "□ 雇用契約を締結する\n"
        "□ 研修の受講\n"
        "□ 研究データ等の情報へのアクセス権限\n"
        "□ サイバー攻撃\n"
        "□ その他"
    )

    # Table 22: §4(2) 実行確認 (1r×1c)
    tbl = _add_table(doc, 1, 1)
    tbl.cell(0, 0).text = "□ 実行可能であることを確認しました"

    # Table 23: 同意確認 (1r×1c)
    tbl = _add_table(doc, 1, 1)
    tbl.cell(0, 0).text = (
        "□ 研究体制におけるPI、Co-PI及び研究参画者の同意を得ました\n"
        "□ 研究代表機関及び共同研究機関の担当部署の確認を得ました"
    )

    out = DUMMY_DIR / "r08youshiki_besshi5.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name} ({len(doc.tables)} tables)")


# ============================================================================
# r08youshiki_betten.docx — 別添: 13 tables (one per item)
# ============================================================================

def create_betten():
    doc = Document()

    # Header paragraph with 氏名
    doc.add_paragraph("氏　名　＿＿＿＿＿＿（＿＿＿・＿＿＿・＿＿＿）")

    # Tables 0-12: 13 items (same structure as besshi5 items)
    table_specs = [
        (3, 2, False),   # ① education_history
        (3, 2, False),   # ② career_history
        (3, 5, True),    # ③ funding_history
        (3, 5, True),    # ④ non_research_support
        (3, 5, True),    # ⑤ publications
        (3, 5, True),    # ⑥ patents
        (3, 3, True),    # ⑦ foreign_talent_programs
        (3, 2, True),    # ⑧ disciplinary_history
        (3, 2, True),    # ⑨ list_status
        (3, 2, True),    # ⑩ listed_entity_affiliation
        (3, 3, True),    # ⑪ listed_entity_relationships
        (3, 2, True),    # ⑫ residency_status
        (1, 2, False),   # ⑬ nationality
    ]

    for rows, cols, has_header in table_specs:
        total_rows = rows + (1 if has_header else 0)
        tbl = _add_table(doc, total_rows, cols)
        if has_header:
            tbl.cell(0, 0).text = "ヘッダー"

    out = DUMMY_DIR / "r08youshiki_betten.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name} ({len(doc.tables)} tables)")


# ============================================================================
# r08youshiki6.xlsx — 様式6: 申請概要
# ============================================================================

def create_youshiki6():
    wb = openpyxl.Workbook()

    # Sheet: 様式6
    ws = wb.active
    ws.title = "様式6"

    # Header row (row 20)
    headers = {
        4: "研究テーマ", 5: "キーワード", 6: "研究分野", 7: "タイプ",
        8: "重複応募", 9: "研究課題名", 10: "役職", 11: "代表者名",
        12: "代表機関", 13: "分担機関", 14: "研究期間",
        15: "1年目", 16: "2年目", 17: "3年目", 18: "4年目", 19: "5年目",
        20: "総額", 21: "メール", 22: "郵送先", 23: "種別",
        24: "中小企業", 25: "スタートアップ", 26: "大学発SU",
        27: "分担中小", 28: "分担SU", 29: "分担大学発SU",
    }
    for col, header in headers.items():
        ws.cell(row=20, column=col).value = header

    # Data row 21 (target row for fill_youshiki6)
    # T21 = SUM formula
    ws.cell(row=21, column=20).value = "=SUM(O21:S21)"

    # Sheet: リスト (dropdown source)
    ws_list = wb.create_sheet("リスト")
    themes = [
        "(1) 知能に関する研究",
        "(23) 医療・医工学に関する基礎研究",
    ]
    for i, theme in enumerate(themes):
        ws_list.cell(row=4 + i, column=1).value = theme

    out = DUMMY_DIR / "r08youshiki6.xlsx"
    wb.save(str(out))
    print(f"  ✓ {out.name}")


# ============================================================================
# r08youshiki7.xlsx — 様式7: 研究者一覧
# ============================================================================

def create_youshiki7():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "採択課題抜粋"

    # Row 23: header
    ws.cell(row=23, column=2).value = "研究課題名"
    ws.cell(row=23, column=3).value = "研究機関名"
    ws.cell(row=23, column=4).value = "氏名"
    ws.cell(row=23, column=5).value = "部局・職"

    # Rows 24+ are data (fill_youshiki7 writes here)

    out = DUMMY_DIR / "r08youshiki7.xlsx"
    wb.save(str(out))
    print(f"  ✓ {out.name}")


# ============================================================================
# r08youshiki8.xlsx — 様式8: 連絡先
# ============================================================================

def create_youshiki8():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    ws.cell(row=2, column=2).value = "メールアドレス"
    # Rows 3-12: email data slots

    out = DUMMY_DIR / "r08youshiki8.xlsx"
    wb.save(str(out))
    print(f"  ✓ {out.name}")


# ============================================================================
# Narrative stubs — inject_narrative.py のE2Eテスト用
# ============================================================================

def create_narrative_stubs():
    """Create minimal narrative docx stubs for inject_narrative.py testing.

    Each stub contains: heading, body paragraph, and a table.
    """
    for name, heading, body_text in [
        ("youshiki1_2_narrative.docx",
         "1. 研究目的（ダミー）",
         "これはE2Eテスト用のダミー本文です。様式1-2の研究計画詳細がここに挿入されます。"),
        ("youshiki1_3_narrative.docx",
         "(1) 研究概要（ダミー）",
         "これはE2Eテスト用のダミー本文です。様式1-3の追加説明事項がここに挿入されます。"),
    ]:
        doc = Document()
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body_text)
        tbl = doc.add_table(rows=2, cols=3)
        tbl.cell(0, 0).text = "項目"
        tbl.cell(0, 1).text = "内容"
        tbl.cell(0, 2).text = "備考"
        tbl.cell(1, 0).text = "ダミーデータ"
        tbl.cell(1, 1).text = "テスト用"
        tbl.cell(1, 2).text = "inject_narrative.py 検証"

        out = DUMMY_DIR / name
        doc.save(str(out))
        print(f"  ✓ {out.name}")


# ============================================================================
# Main
# ============================================================================

def main():
    print("=== E2Eテスト用スタブファイル生成 ===")
    print()
    create_youshiki1_5()
    create_besshi5()
    create_betten()
    create_narrative_stubs()
    create_youshiki6()
    create_youshiki7()
    create_youshiki8()
    print()
    print("✓ 全スタブファイルを生成しました")


if __name__ == "__main__":
    main()
